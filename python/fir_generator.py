import os
import json
import re
import google.generativeai as genai
from dotenv import load_dotenv
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load API key
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

# Theft categories and specific details that should be in FIRs based on your examples
THEFT_CATEGORIES = {
    "Pickpocketing": [
        "Exact time and location when victim noticed the item was missing",
        "Complete description of stolen items including brand, model, value",
        "Activities before noticing the theft (where were you coming from/going to)",
        "Any CCTV cameras in the area",
        "Any suspects or people who were nearby",
        "IMEI numbers for phones or serial numbers for electronics"
    ],
    "Burglary": [
        "Exact date and time when you discovered the break-in",
        "How entry was made to your property (broken lock, window, etc.)",
        "Complete list of items stolen with approximate values",
        "Who was at home during the incident, if anyone",
        "Details of any security systems or CCTV",
        "When you last saw the items before they were stolen"
    ],
    "ATM Fraud": [
        "Last legitimate transaction details before the fraud",
        "Bank account details and card information (card type, issuing bank)",
        "Transaction IDs, dates, times of fraudulent transactions",
        "When and how you discovered the fraud",
        "Whether you've contacted your bank and what actions they've taken",
        "Details of any SMS alerts received about transactions"
    ],
    "Online Fraud": [
        "Platform/website where fraud occurred",
        "Transaction details (IDs, times, amounts)",
        "Communication received from fraudsters (emails, messages)",
        "How payment was made (credit card, UPI, etc.)",
        "When and how you discovered the fraud",
        "Any actions already taken (bank notification, platform reporting)"
    ],
    "Vehicle Theft": [
        "Make, model, color, registration number of vehicle",
        "Exact location where vehicle was parked with landmarks",
        "Last time you saw the vehicle before theft",
        "Any valuables or documents inside the vehicle",
        "Any security features (steering lock, GPS tracker)",
        "Details of parking receipt or payment if applicable"
    ],
    "Mobile Phone Theft": [
        "Make, model, color, IMEI number of phone",
        "SIM card details including phone number and service provider",
        "Any accounts logged in on the phone",
        "Exact location and circumstances of theft",
        "Whether remote locking/tracking was attempted",
        "Value of the phone and accessories stolen"
    ],
    "Missing Person": [
        "Full name, age, gender, appearance details of missing person",
        "Clothing and accessories they were wearing when last seen",
        "Last known location and time they were seen",
        "Any health conditions or medications needed",
        "Any reason they might have left or places they might go",
        "Their mobile number and whether it's reachable"
    ],
    "Identity Theft": [
        "Which documents or ID proofs were compromised",
        "Any unauthorized accounts opened in your name",
        "Transactions or activities conducted using your identity",
        "When and how you discovered the identity theft",
        "Whether you've notified relevant authorities (bank, passport office)",
        "Any financial loss incurred"
    ]
}

# IPC Sections for different theft types
IPC_SECTIONS = {
    "Pickpocketing": ["IPC", "Section 379 (Theft)"],
    "Burglary": ["IPC", "Section 454 (House-breaking)", "Section 457 (House-breaking by night)"],
    "ATM Fraud": ["IPC", "Section 420 (Cheating)", "IT Act", "Section 66D (Cheating by personation using computer resource)"],
    "Online Fraud": ["IPC", "Section 420 (Cheating)", "IT Act", "Section 66D (Cheating by personation using computer resource)"],
    "Vehicle Theft": ["IPC", "Section 379 (Theft)", "Section 381 (Theft by clerk or servant of property in possession of master)"],
    "Mobile Phone Theft": ["IPC", "Section 379 (Theft)"],
    "Missing Person": ["CrPC", "Section 174 (Police to enquire and report on missing person)"],
    "Identity Theft": ["IPC", "Section 420 (Cheating)", "IT Act", "Section 66C (Identity theft)"]
}

def extract_details(user_complaint):
    """Extract structured information from theft complaints using Gemini API"""
    prompt = f"""You are an AI assistant that extracts structured information from theft complaints for a First Information Report (FIR).
Extract the following details in JSON format:
- type_of_theft (Must be one of: Pickpocketing, Burglary, ATM Fraud, Online Fraud, Vehicle Theft, Mobile Phone Theft, Missing Person, Identity Theft)
- date_time (when the incident occurred)
- location (specific place where the incident occurred)
- items_stolen (list all items stolen) or missing_person_details (if applicable)
- estimated_value (total value of stolen items/money in INR)
- mode_of_theft (how the theft was committed)
- suspect_description (if available)
- witness_details (if available)
- complainant_name (name of the person reporting)
- complainant_contact (phone number if mentioned)
- complainant_address (address if mentioned)
- complainant_id_details (any ID information mentioned)
- additional_details (any other relevant information)

Example Output:
{{
  "type_of_theft": "Pickpocketing",
  "date_time": "October 15, 2024 around 3:30 PM",
  "location": "Central Mall food court, MG Road, Mumbai",
  "items_stolen": "Wallet containing INR 5,000, driver's license, 2 credit cards, office ID",
  "estimated_value": "8,500",
  "mode_of_theft": "Distraction while victim was at food court",
  "suspect_description": "Young man wearing black t-shirt",
  "witness_details": "Food court staff noticed suspicious person",
  "complainant_name": "Not specified",
  "complainant_contact": "Not specified",
  "complainant_address": "Not specified",
  "complainant_id_details": "Driver's license was stolen",
  "additional_details": "CCTV cameras in the area weren't working properly. Unauthorized transaction of INR 3,500 made at electronics store in the same mall."
}}

Now, extract details from this complaint:
\"\"\"{user_complaint}\"\"\""""

    try:
        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)
        
        # Extract text response
        if response and hasattr(response, 'candidates') and response.candidates:
            response_text = response.candidates[0].content.parts[0].text.strip()
            
            # Find JSON in the response (in case there's additional text)
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
                response_dict = json.loads(json_str)
                return response_dict
            else:
                return {"error": "Could not find valid JSON in the response"}
        else:
            return {"error": "No valid response from API"}

    except json.JSONDecodeError as e:
        return {"error": f"JSON parsing failed: {str(e)}"}
    except Exception as e:
        return {"error": f"Error processing complaint: {str(e)}"}

def identify_missing_details(extracted_data):
    """Identify missing or incomplete details that would be needed for a complete FIR"""
    
    # First, get the mandatory fields for any FIR
    mandatory_fields = {
        "complainant_name": "What is your full name?",
        "complainant_contact": "What is your contact number?",
        "complainant_address": "What is your complete residential address?",
        "date_time": "When exactly did the incident occur (date and time)?",
        "location": "What is the exact location where the incident occurred (with landmarks if possible)?"
    }
    
    missing_fields = {}
    
    # Check for mandatory fields
    for field, question in mandatory_fields.items():
        if field not in extracted_data or not extracted_data[field] or extracted_data[field].lower() == "not specified":
            missing_fields[field] = question
    
    # Get theft type specific missing details
    theft_type = extracted_data.get("type_of_theft", "Unknown")
    
    if theft_type in THEFT_CATEGORIES:
        # Get the specific details needed for this type of theft
        required_details = THEFT_CATEGORIES[theft_type]
        
        # Use Gemini to analyze what specific details might be missing
        prompt = f"""
        I have a complaint about a {theft_type} incident. Based on the extracted information below, identify which specific critical details are missing that police would need for a complete FIR.

        Extracted information:
        {json.dumps(extracted_data, indent=2)}
        
        For this type of incident, police typically need these specific details:
        {json.dumps(required_details, indent=2)}
        
        Identify the top 3-5 most important specific details that are missing or need clarification from this account. For each missing detail, provide a clear question that a police officer would ask to get this information.
        
        Format your response as a JSON array of objects with 'detail_missing' and 'question' fields. Example:
        [
            {{"detail_missing": "IMEI number of stolen phone", "question": "What is the IMEI number of your stolen phone? It can usually be found on the box or by dialing *#06# on a similar phone."}},
            {{"detail_missing": "Exact time of theft", "question": "At what exact time did you notice your phone was missing?"}}
        ]
        
        Only include truly missing information, not details that are already provided.
        """
        
        try:
            model = genai.GenerativeModel("gemini-1.5-pro")
            response = model.generate_content(prompt)
            
            # Extract text response
            if response and hasattr(response, 'candidates') and response.candidates:
                response_text = response.candidates[0].content.parts[0].text.strip()
                
                # Find JSON in the response
                json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
                if json_match:
                    json_str = json_match.group(0)
                    missing_details = json.loads(json_str)
                    
                    # Add these to our missing fields
                    for item in missing_details:
                        field_name = item["detail_missing"].lower().replace(" ", "_")
                        missing_fields[field_name] = item["question"]
                    
        except Exception as e:
            print(f"Warning: Error analyzing missing details: {str(e)}")
            # If Gemini fails, fall back to basic checks
            
            # Check for stolen items and value if applicable
            if theft_type != "Missing Person":
                if "items_stolen" not in extracted_data or not extracted_data["items_stolen"] or extracted_data["items_stolen"].lower() == "not specified":
                    missing_fields["items_stolen"] = "What items were stolen exactly?"
                
                if "estimated_value" not in extracted_data or not extracted_data["estimated_value"] or extracted_data["estimated_value"].lower() == "not specified":
                    missing_fields["estimated_value"] = "What is the approximate value of the stolen items (in Rs.)?"
            else:
                if "missing_person_details" not in extracted_data or not extracted_data["missing_person_details"] or extracted_data["missing_person_details"].lower() == "not specified":
                    missing_fields["missing_person_details"] = "Can you provide details about the missing person (name, age, appearance, what they were wearing)?"
    
    return missing_fields

def generate_detailed_questions(extracted_data):
    """Generate follow-up questions to get a comprehensive account for the FIR"""
    theft_type = extracted_data.get("type_of_theft", "Unknown")
    
    # Use Gemini to generate contextual follow-up questions based on the specific case
    prompt = f"""
    You are an experienced police officer taking a First Information Report (FIR) for a {theft_type} case. 
    
    Here is the information you already have from the complainant:
    {json.dumps(extracted_data, indent=2)}
    
    Based on the above information and your experience with similar cases, generate 5-8 specific follow-up questions that would help:
    1. Establish a clear timeline of events
    2. Gather evidence that could help solve the case
    3. Get specific details needed for a complete legal FIR document
    4. Get details that would differentiate this case from similar ones

    IMPORTANT: DO NOT ask for the following basic information that would be collected separately:
    - Complainant's name
    - Complainant's contact number
    - Complainant's address
    - Complainant's ID details
    
    
    Format your response as a JSON array of strings, each containing one question.
    Questions should be direct and specific to this case, not generic.
    """
    
    try:
        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)
        
        # Extract text response
        if response and hasattr(response, 'candidates') and response.candidates:
            response_text = response.candidates[0].content.parts[0].text.strip()
            
            # Find JSON in the response
            json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
                questions = json.loads(json_str)
                return questions
            else:
                return []
        else:
            return []

    except Exception as e:
        print(f"Warning: Error generating follow-up questions: {str(e)}")
        return []

def generate_legal_narrative(extracted_data, additional_info):
    """Generate a professionally formatted legal narrative for the FIR"""
    # Combine extracted data with additional info
    complaint_data = {**extracted_data, **additional_info}
    
    prompt = f"""
    Generate a professionally formatted legal narrative for a First Information Report (FIR) based on the following theft incident:
    
    {json.dumps(complaint_data, indent=2)}
    
    The narrative should:
    1. Use formal legal language appropriate for a police document
    2. Begin with details of the complainant (name, address, occupation)
    3. Describe the incident chronologically with precise details
    4. Mention exact time, date and location
    5. Include specific details about stolen items and their value (or person missing, if applicable)
    6. Reference any witnesses or evidence (if available)
    7. Include suspect descriptions (if available)
    8. End with a formal declaration that the complainant believes the information to be true
    
    Format as a single, comprehensive paragraph without headings or bullet points.
    Use precise, unambiguous language as would be found in an official police report.
    The content should resemble the level of detail found in actual FIRs, similar to these examples from Indian police stations.
    """

    try:
        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)
        
        if hasattr(response, 'text'):
            return response.text.strip()
        elif hasattr(response, 'candidates') and response.candidates:
            return response.candidates[0].content.parts[0].text.strip()
        else:
            return "Error: Unable to generate FIR narrative content."
    except Exception as e:
        return f"Error generating FIR narrative: {str(e)}"

def create_fir_document(narrative, complaint_data, output_filename="FIR_Document.docx"):
    """Create an editable Word document containing the FIR information"""
    try:
        # Create a new Document
        doc = Document()
        
        # Set page margin
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('FIRST INFORMATION REPORT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add bilingual title (English and Hindi)
        subtitle = doc.add_paragraph()
        subtitle.add_run('First Information contents (प्रथम खबर हकिगत):').bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add current date
        current_date = datetime.now().strftime("%d/%m/%Y")
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f"दिनांक: {current_date}").bold = True
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add FIR number (placeholder)
        header_table = doc.add_table(rows=1, cols=2)
        header_table.style = 'Table Grid'
        header_table.columns[0].width = Inches(3)
        header_table.columns[1].width = Inches(3)
        
        cells = header_table.rows[0].cells
        cells[0].text = "FIR No.: _________________"
        cells[1].text = "Police Station: _________________"
        
        # Add a blank line
        doc.add_paragraph()
        
        # Add complainant information
        complainant_name = complaint_data.get('complainant_name', 'Not specified')
        complainant_contact = complaint_data.get('complainant_contact', 'Not specified')
        complainant_address = complaint_data.get('complainant_address', 'Not specified')
        
        complainant_para = doc.add_paragraph()
        complainant_para.add_run(f"{complainant_name}, वय- ").bold = True
        complainant_para.add_run("___ वर्ष, ")
        complainant_para.add_run("व्यवसाय- ").bold = True
        complainant_para.add_run("______________, ")
        complainant_para.add_run("पत्ता- ").bold = True
        complainant_para.add_run(f"{complainant_address}, ")
        complainant_para.add_run("मो.नं.- ").bold = True
        complainant_para.add_run(f"{complainant_contact}.")
        
        # Add a blank line
        doc.add_paragraph()
        
        # Add the narrative with proper paragraph formatting
        narrative_paragraph = doc.add_paragraph(narrative)
        
        # Add a blank line
        doc.add_paragraph()
        
        # Space for signature
        doc.add_paragraph("\n\n")
        sign_paragraph = doc.add_paragraph("तरी वरील जबाब मी सांगितल्याप्रमाणे बरोबर व खरा आहे तो मला वाचून दाखविला असून तो माझे सांगणे प्रमाणे बरोबर व खरा आहे.")
        sign_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph("\n\n")
        sign_table = doc.add_table(rows=1, cols=2)
        sign_table.style = 'Table Grid'
        
        cells = sign_table.rows[0].cells
        cells[0].text = "जबाब लिहून घेणाऱ्याची सही\n\n\n\n________________"
        cells[1].text = "जबाब देणाऱ्याची सही\n\n\n\n________________"
        
        # Save the document
        doc.save(output_filename)
        return output_filename
    
    except Exception as e:
        return f"Error creating Word document: {str(e)}"

def determine_applicable_sections(theft_type):
    """Return applicable IPC sections for the theft type"""
    sections = IPC_SECTIONS.get(theft_type, [])
    if not sections:
        return "To be determined by investigating officer"
    
    result = ""
    for i in range(0, len(sections), 2):
        if i < len(sections):
            result += sections[i]
            if i+1 < len(sections):
                result += ", " + sections[i+1]
            result += "; "
    
    return result.strip("; ")

def main():
    """Main function to generate the FIR First Information Content"""
    # Check if python-docx is installed, if not inform the user to install it
    try:
        import docx
    except ImportError:
        print("ERROR: The 'python-docx' package is required but not installed.")
        print("Please install it by running: pip install python-docx")
        return
    
    print("=== FIR Document Generator ===")
    print("Please provide details about the incident:")
    
    # Get user complaint
    user_complaint = input("Describe the incident in detail:\n")
    
   # Process the complaint
    print("\nAnalyzing your complaint...\n")
    extracted_data = extract_details(user_complaint)
    
    if "error" in extracted_data:
        print(f"Error: {extracted_data['error']}")
        return
    
    theft_type = extracted_data.get('type_of_theft', 'Unknown')
    print(f"Incident type identified: {theft_type}")
    print(f"Applicable sections: {determine_applicable_sections(theft_type)}")
    
    # First get missing mandatory fields
    missing_fields = identify_missing_details(extracted_data)
    additional_info = {}
    
    # Then generate follow-up questions
    follow_up_questions = generate_detailed_questions(extracted_data)
    
    # Create a unified set of questions without duplicates
    all_questions = {}
    
    # Add missing fields questions
    for field, question in missing_fields.items():
        all_questions[question] = field
    
    # Add follow-up questions, avoiding duplicates based on similarity
    question_count = 0
    for question in follow_up_questions:
        # Check if this question is too similar to an existing one
        is_duplicate = False
        for existing_question in all_questions.keys():
            # Simple check for similarity - if questions share multiple words
            common_words = set(question.lower().split()) & set(existing_question.lower().split())
            if len(common_words) >= 3:  # If 3+ words match, consider it similar
                is_duplicate = True
                break
        
        if not is_duplicate:
            field_name = f"detail_{question_count}"
            all_questions[question] = field_name
            question_count += 1
    
    # Now ask all questions without repetition
    if all_questions:
        print("\nPlease provide the following information for your FIR:")
        for question, field in all_questions.items():
            response = input(f"{question} ")
            additional_info[field] = response 
    # Generate narrative for FIR
    print("\nGenerating FIR narrative...")
    narrative = generate_legal_narrative(extracted_data, additional_info)
    
    # Create the Word document
    output_file = create_fir_document(narrative, {**extracted_data, **additional_info})
    
    if output_file.startswith("Error"):
        print(f"Error: {output_file}")
    else:
        print("\nFIR Document Summary:")
        print("-" * 50)
        print(f"Complainant: {extracted_data.get('complainant_name', additional_info.get('complainant_name', 'Not specified'))}")
        print(f"Incident Type: {theft_type}")
        print(f"Date/Time: {extracted_data.get('date_time', additional_info.get('date_time', 'Not specified'))}")
        print(f"Location: {extracted_data.get('location', additional_info.get('location', 'Not specified'))}")
        print("-" * 50)
        print(f"\nFIR document has been saved to: {output_file}")
        print("You can now open, edit, and print this document as needed.")
        
        # Ask if user wants to open the document now (Windows only)
        if os.name == 'nt':  # Windows
            open_doc = input("\nWould you like to open the document now? (y/n): ")
            if open_doc.lower() == 'y':
                os.startfile(output_file)
        else:
            print("\nYou can open the document using any word processor like Microsoft Word, LibreOffice Writer, etc.")

if __name__ == "__main__":
    main()